"""
FastAPI backend pour SAIM Conseil RAG.
Expose /api/chat (streaming) et /api/status.
Fallback : Groq → Ollama local si quota dépassé.
"""

import os
import re
import json
import aiohttp
from dotenv import load_dotenv

load_dotenv()

from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
from typing import Optional
from groq import AsyncGroq
from rag import build_messages_with_rag, is_db_ready

_api_key = os.environ.get("GROQ_API_KEY")
if not _api_key:
    raise RuntimeError("GROQ_API_KEY manquant — vérifiez backend/.env")

groq_client = AsyncGroq(api_key=_api_key)

GROQ_MODEL = "llama-3.3-70b-versatile"
OLLAMA_URL = "http://localhost:11434/api/chat"
OLLAMA_MODELS = ["qwen3:8b", "llama3.1:latest"]  # fallback local gratuit

app = FastAPI(title="SAIM Conseil RAG Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_methods=["*"],
    allow_headers=["*"],
)


class ChatRequest(BaseModel):
    messages: list[dict]


class InvoiceRequest(BaseModel):
    prompt: str
    image_base64: Optional[str] = None
    mode: str = "facture"  # "facture" | "devis"


class MarketingRequest(BaseModel):
    mode: str  # "analyse" | "planning" | "stratégie"
    data: dict


class CommercialRequest(BaseModel):
    mode: str   # "analyse" | "proposition"
    prompt: str


class ProjectRequest(BaseModel):
    mode: str   # "planning" | "risks"
    prompt: str


async def groq_json_with_fallback(messages: list, max_tokens: int = 3500) -> dict:
    """Appel Groq JSON mode + fallback Ollama si rate limit."""
    try:
        completion = await groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=messages,
            max_tokens=max_tokens,
            temperature=0.3,
            stream=False,
            response_format={"type": "json_object"},
        )
        return json.loads(completion.choices[0].message.content)
    except Exception as e:
        if "rate_limit" not in str(e) and "429" not in str(e):
            raise
        # Fallback Ollama
        ollama_msgs = list(messages)
        if ollama_msgs and ollama_msgs[-1]["role"] == "user":
            ollama_msgs[-1] = {
                **ollama_msgs[-1],
                "content": ollama_msgs[-1]["content"] + "\n\nRetourne UNIQUEMENT du JSON valide, sans aucun texte avant ou après.",
            }
        for model in OLLAMA_MODELS:
            try:
                async with aiohttp.ClientSession() as session:
                    payload = {"model": model, "messages": ollama_msgs, "stream": False}
                    async with session.post(OLLAMA_URL, json=payload, timeout=aiohttp.ClientTimeout(total=180)) as resp:
                        if resp.status != 200:
                            continue
                        data = await resp.json()
                        content = data.get("message", {}).get("content", "")
                        match = re.search(r'\{[\s\S]*\}', content)
                        if match:
                            return json.loads(match.group())
            except Exception:
                continue
        raise Exception("Groq quota dépassé et Ollama indisponible")


async def stream_groq(messages: list):
    """Génère les tokens via Groq (async). Lève une exception si rate limit."""
    completion = await groq_client.chat.completions.create(
        model=GROQ_MODEL,
        messages=messages,
        temperature=0.1,
        max_tokens=1024,
        top_p=0.9,
        stream=True,
    )
    async for chunk in completion:
        content = chunk.choices[0].delta.content
        if content:
            yield content


async def stream_ollama(messages: list):
    """Génère les tokens via Ollama local (fallback gratuit)."""
    for model in OLLAMA_MODELS:
        try:
            async with aiohttp.ClientSession() as session:
                payload = {"model": model, "messages": messages, "stream": True}
                async with session.post(OLLAMA_URL, json=payload, timeout=aiohttp.ClientTimeout(total=120)) as resp:
                    if resp.status != 200:
                        continue
                    async for line in resp.content:
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            token = data.get("message", {}).get("content", "")
                            if token:
                                yield token
                            if data.get("done"):
                                return
                        except json.JSONDecodeError:
                            continue
            return  # succès
        except Exception:
            continue
    raise Exception("Ollama indisponible.")


async def generate(messages: list):
    """Essaie Groq, bascule sur Ollama si quota dépassé."""
    try:
        async for token in stream_groq(messages):
            yield token
    except Exception as e:
        err_str = str(e)
        if "rate_limit_exceeded" in err_str or "429" in err_str:
            yield "*(quota Groq atteint — bascule sur le modèle local…)*\n\n"
            async for token in stream_ollama(messages):
                yield token
        else:
            print(f"[Groq non-rate-limit error] {type(e).__name__}: {e}")
            raise


@app.get("/api/status")
def status():
    return {
        "status": "ok",
        "rag_ready": is_db_ready(),
        "model": GROQ_MODEL,
        "fallback": OLLAMA_MODELS[0],
    }


@app.post("/api/invoice/extract")
async def extract_invoice(req: InvoiceRequest):
    """Extrait les données de facturation depuis un prompt texte (+ image optionnelle)."""
    from datetime import date, timedelta
    import random
    today_str = date.today().strftime("%d/%m/%Y")
    due_str = (date.today() + timedelta(days=30)).strftime("%d/%m/%Y")
    year = date.today().year
    num = random.randint(1, 999)

    is_devis = req.mode == "devis"
    prefix = "DEVIS" if is_devis else "FACT"
    valid_str = (date.today() + timedelta(days=30)).strftime("%d/%m/%Y")

    if is_devis:
        system = (
            f"Tu es un commercial expert PME camerounaise. Aujourd'hui : {today_str}.\n"
            "Génère un devis professionnel complet depuis la description de l'utilisateur.\n"
            "Retourne UNIQUEMENT un JSON avec ces champs camelCase :\n"
            f"- invoiceNumber : \"DEVIS-{year}-{num:03d}\"\n"
            f"- date : \"{today_str}\"\n"
            f"- dueDate : \"{valid_str}\" (date limite de validité)\n"
            f"- validUntil : \"{valid_str}\"\n"
            "- client : {{name, address, email, phone}}\n"
            "- items : [{{description, qty, unitPrice, total}}] — un objet par prestation/produit\n"
            "- subtotal, tvaRate (19.25 si applicable), tva, total\n"
            "- currency : \"FCFA\"\n"
            "- payment_terms : délai de paiement (ex: \"30 jours\")\n"
            "- payment_methods : liste des moyens acceptés (ex: [\"Mobile Money MTN\", \"Orange Money\", \"Virement\"])\n"
            "- conditions : conditions générales courtes\n"
            "- notes : remarques ou commentaires\n"
            "Calcule TVA 19,25% sauf si le client est exonéré. Sois professionnel et complet."
        )
    else:
        system = (
            f"Tu es un expert en facturation camerounaise. La date d'aujourd'hui est {today_str}.\n\n"
            "À partir du message de l'utilisateur, génère un objet JSON de facture.\n"
            "Le JSON doit contenir EXACTEMENT ces champs (noms camelCase stricts) :\n"
            "- invoiceNumber : string ex \"FACT-" + str(year) + f"-{num:03d}\"\n"
            f"- date : \"{today_str}\"\n"
            f"- dueDate : \"{due_str}\"\n"
            "- client : objet {name, address, email, phone} — extraire le nom du client du message\n"
            "- items : tableau d'objets {description, qty, unitPrice, total} — un objet par article mentionné\n"
            "  * qty = quantité (entier)\n"
            "  * unitPrice = prix unitaire (nombre)\n"
            "  * total = qty * unitPrice (calculé)\n"
            "- subtotal : somme de tous les items.total\n"
            "- tvaRate : 19.25\n"
            "- tva : round(subtotal * 0.1925, 0)\n"
            "- total : subtotal + tva\n"
            "- currency : \"FCFA\"\n"
            "- notes : string vide ou remarques éventuelles\n\n"
            "Si des informations manquent (adresse, email…) laisse le champ à chaîne vide.\n"
            "Ne demande rien, génère la facture avec ce qui est fourni."
        )

    # Choix du modèle et format du message selon présence d'image
    if req.image_base64:
        # Modèle vision — format multimodal, pas de json_object mode
        vision_model = "llama-3.2-11b-vision-preview"
        user_message = {
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {
                        # req.image_base64 contient le data URL complet (data:image/png;base64,...)
                        "url": req.image_base64 if req.image_base64.startswith("data:") else f"data:image/jpeg;base64,{req.image_base64}"
                    },
                },
                {
                    "type": "text",
                    "text": (
                        system + "\n\nAnalyse cette image et extrais toutes les informations "
                        "de facturation visibles (articles, quantités, prix, client…). "
                        + (f"\nNote additionnelle de l'utilisateur : {req.prompt}" if req.prompt.strip() else "")
                        + "\n\nRetourne UNIQUEMENT le JSON brut sans markdown."
                    ),
                },
            ],
        }
        messages = [user_message]
        use_json_mode = False
        model = vision_model
    else:
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": req.prompt},
        ]
        use_json_mode = True
        model = GROQ_MODEL

    raw_content = ""
    try:
        kwargs: dict = dict(
            model=model,
            messages=messages,
            max_tokens=2048,
            temperature=0.1,
            stream=False,
        )
        if use_json_mode:
            kwargs["response_format"] = {"type": "json_object"}

        completion = await groq_client.chat.completions.create(**kwargs)
        raw_content = completion.choices[0].message.content.strip()
        # Pour le mode vision, extraire le JSON depuis le texte (peut contenir du markdown)
        if not use_json_mode:
            match = re.search(r'\{[\s\S]*\}', raw_content)
            if not match:
                return JSONResponse(status_code=422, content={
                    "error": "Le modèle vision n'a pas retourné de JSON exploitable.",
                    "raw": raw_content[:500]
                })
            raw_content = match.group()
        invoice_data = json.loads(raw_content)

        # Recalcul côté serveur pour garantir la cohérence des totaux
        items = invoice_data.get("items", [])
        for it in items:
            it["qty"] = int(it.get("qty") or 0)
            it["unitPrice"] = float(it.get("unitPrice") or 0)
            it["total"] = round(it["qty"] * it["unitPrice"], 2)
        subtotal = round(sum(it["total"] for it in items), 2)
        tva = round(subtotal * 0.1925, 0)
        invoice_data["items"] = items
        invoice_data["subtotal"] = subtotal
        invoice_data["tvaRate"] = 19.25
        invoice_data["tva"] = tva
        invoice_data["total"] = subtotal + tva
        invoice_data["currency"] = "FCFA"

        return JSONResponse(content=invoice_data)
    except json.JSONDecodeError:
        print(f"[INVOICE JSON ERROR] raw=\n{raw_content}")
        return JSONResponse(status_code=422, content={"error": f"Réponse non-JSON du modèle : {raw_content[:300]}"})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/marketing")
async def marketing(req: MarketingRequest):
    from datetime import date, timedelta
    today = date.today()

    if req.mode == "analyse":
        system = """Tu es un expert en marketing, segmentation client et psychologie du consommateur africain.
Génère exactement 3 profils clients archétypes pour le produit/service décrit.
Retourne UNIQUEMENT un JSON valide avec cette structure exacte:
{
  "profiles": [
    {
      "archetype": "Nom de l'archétype (ex: L'Entrepreneur Ambitieux)",
      "emoji": "emoji représentatif du profil",
      "age_range": "ex: 28-38",
      "gender": "Homme / Femme / Mixte",
      "location": "villes au Cameroun",
      "income": "Faible / Moyen / Élevé",
      "income_score": 65,
      "interests": ["intérêt1", "intérêt2", "intérêt3"],
      "pain_points": ["problème1", "problème2", "problème3"],
      "buying_behavior": "description détaillée du comportement d'achat",
      "preferred_channels": ["Instagram", "WhatsApp"],
      "key_message": "Message marketing percutant pour ce profil (1 phrase)",
      "conversion_tips": ["conseil1", "conseil2"]
    }
  ]
}"""
        user_content = (
            f"Produit/Service: {req.data.get('service', '')}\n"
            f"Secteur: {req.data.get('sector', '')}\n"
            f"Prix: {req.data.get('price', '')}\n"
            f"Description: {req.data.get('description', '')}\n"
            f"Zone géographique cible: {req.data.get('zone', 'Cameroun')}"
        )

    elif req.mode == "planning":
        system = f"""Tu es un expert en stratégie de contenu et marketing digital pour les entreprises africaines.
Génère un calendrier éditorial sur 4 semaines à partir du {today.strftime("%d/%m/%Y")}.
Retourne UNIQUEMENT un JSON valide:
{{
  "title": "Titre du calendrier éditorial",
  "calendar": [
    {{
      "date": "JJ/MM/AAAA",
      "day": "Lundi",
      "week": 1,
      "theme": "Thématique de la semaine",
      "content": "Description précise du contenu à créer",
      "type": "Post Instagram",
      "objective": "Notoriété",
      "hashtags": ["#tag1", "#tag2"],
      "cta": "Call-to-action suggéré"
    }}
  ]
}}
Types possibles: Post Instagram, Reels, Story, Vidéo YouTube, Article Blog, Email, LinkedIn, WhatsApp, TikTok.
Objectifs possibles: Notoriété, Engagement, Conversion, Fidélisation, Trafic.
Génère 3 à 4 publications par semaine, varie les types et objectifs."""
        user_content = (
            f"Produit/Service: {req.data.get('service', '')}\n"
            f"Secteur: {req.data.get('sector', '')}\n"
            f"Cible: {req.data.get('target', '')}\n"
            f"Canaux principaux: {req.data.get('channels', '')}\n"
            f"Objectif principal de la campagne: {req.data.get('objective', '')}\n"
            f"Ton de communication: {req.data.get('tone', 'Professionnel et accessible')}"
        )
    elif req.mode == "stratégie":
        # Système de dialogue : vérifier si on a assez d'infos, sinon poser des questions
        prompt = req.data.get('prompt', '').strip()
        duration = req.data.get('duration', '').strip()
        objective = req.data.get('objective', '').strip()
        
        # Si infos manquent, retourner des questions
        if not duration or not objective:
            questions = []
            if not duration:
                questions.append({
                    "id": "duration",
                    "question": "Quelle est la durée de la stratégie marketing souhaitée ?",
                    "type": "select",
                    "options": ["3 mois", "6 mois", "1 an", "2 ans", "Personnalisé"]
                })
            if not objective:
                questions.append({
                    "id": "objective", 
                    "question": "Quel est l'objectif principal de cette stratégie ?",
                    "type": "select",
                    "options": ["Augmenter les ventes", "Accroître la notoriété", "Fidéliser les clients", "Lancer un nouveau produit", "Personnalisé"]
                })
            
            return JSONResponse(content={
                "type": "questions",
                "questions": questions,
                "message": "J'ai besoin de plus d'informations pour créer une stratégie marketing complète."
            })
        
        # Si on a toutes les infos, générer la stratégie
        system = f"""Tu es un expert en stratégie marketing digitale pour les PME camerounaises.
Crée une stratégie marketing complète et actionnable sur {duration}.
Retourne UNIQUEMENT un JSON valide avec cette structure :
{{
  "title": "Titre de la stratégie",
  "duration": "{duration}",
  "objective": "{objective}",
  "budget_estimate": "Estimation budgétaire en FCFA",
  "target_audience": "Description détaillée de la cible",
  "competitive_analysis": "Analyse concurrentielle brève",
  "pillars": [
    {{
      "name": "Pilier stratégique 1",
      "description": "Description détaillée",
      "actions": ["Action 1", "Action 2", "Action 3"],
      "kpis": ["KPI 1", "KPI 2"],
      "timeline": "Quand exécuter"
    }}
  ],
  "channels": ["Canal 1", "Canal 2", "Canal 3"],
  "content_strategy": "Stratégie de contenu détaillée",
  "implementation_steps": ["Étape 1", "Étape 2", "Étape 3"],
  "risks_mitigation": "Risques et mesures correctives",
  "success_metrics": ["Métrique 1", "Métrique 2", "Métrique 3"]
}}
Sois concret, chiffré, et adapté au marché camerounais."""
        
        user_content = f"Contexte : {prompt}\nDurée : {duration}\nObjectif : {objective}"
    else:
        return JSONResponse(status_code=400, content={"error": f"Mode inconnu: {req.mode}"})

    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user_content},
    ]

    try:
        completion = await groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=messages,
            max_tokens=4096,
            temperature=0.4,
            stream=False,
            response_format={"type": "json_object"},
        )
        raw = completion.choices[0].message.content.strip()
        data = json.loads(raw)
        return JSONResponse(content=data)
    except json.JSONDecodeError:
        print(f"[MARKETING JSON ERROR] raw=\n{raw}")
        return JSONResponse(status_code=422, content={"error": f"JSON invalide: {raw[:300]}"})
    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/marketing/export")
async def export_marketing_strategy(req: dict):
    """Exporte une stratégie marketing en document Word."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        import io
        
        strategy_data = req.get("strategy", {})
        
        doc = Document()
        doc.add_heading(strategy_data.get("title", "Stratégie Marketing"), 0)
        
        # Ajouter les sections principales
        sections = [
            ("Durée", strategy_data.get("duration", "")),
            ("Objectif", strategy_data.get("objective", "")),
            ("Budget estimé", strategy_data.get("budget_estimate", "")),
            ("Public cible", strategy_data.get("target_audience", "")),
            ("Analyse concurrentielle", strategy_data.get("competitive_analysis", "")),
        ]
        
        for section_title, content in sections:
            if content:
                doc.add_heading(section_title, level=1)
                doc.add_paragraph(content)
        
        # Pilier stratégiques
        if "pillars" in strategy_data:
            doc.add_heading("Piliers Stratégiques", level=1)
            for pillar in strategy_data["pillars"]:
                doc.add_heading(pillar["name"], level=2)
                doc.add_paragraph(pillar["description"])
                
                if "actions" in pillar:
                    doc.add_paragraph("Actions :")
                    for action in pillar["actions"]:
                        doc.add_paragraph(f"• {action}", style='List Bullet')
                
                if "kpis" in pillar:
                    doc.add_paragraph("KPIs :")
                    for kpi in pillar["kpis"]:
                        doc.add_paragraph(f"• {kpi}", style='List Bullet')
        
        # Sauvegarder en mémoire
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=strategie_marketing.docx"}
        )
        
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/commercial")
async def commercial(req: CommercialRequest):
    if req.mode == "analyse":
        system = """Tu es un analyste commercial expert des PME africaines et camerounaises.
Analyse le prospect décrit. Retourne UNIQUEMENT un JSON valide:
{"name":"Nom prospect","sector":"Secteur","company_size":"PME/TPE/Grande entreprise",
"score":7,"score_label":"Fort potentiel","score_reason":"Explication courte",
"purchase_potential":"Description du potentiel d'achat",
"needs":[{"need":"Besoin 1","priority":"Haute","detail":"Explication"}],
"objections":[{"objection":"Objection probable","response":"Comment répondre concrètement"}],
"recommended_approach":"Approche recommandée",
"best_contact":"Meilleur moment/canal","next_steps":["Étape 1","Étape 2","Étape 3"],
"decision_makers":["Qui contacter"],"timeline":"Délai estimé conversion"}
Score de 1 (très faible) à 10 (excellent). Sois concret et adapté au contexte camerounais."""

    elif req.mode == "proposition":
        system = """Tu es un commercial expert du marché camerounais.
Génère une proposition commerciale percutante adaptée aux codes culturels locaux.
Retourne UNIQUEMENT un JSON valide:
{"channel":"WhatsApp/Email/Téléphone/Face à face","prospect_name":"Nom",
"subject":"Objet/titre","hook":"Phrase d'accroche percutante (1 ligne)",
"message":"Message complet prêt à envoyer, formaté pour le canal",
"value_props":["Avantage 1","Avantage 2","Avantage 3"],
"cta":"Appel à l'action précis","follow_up":"Stratégie de relance",
"cultural_notes":"Adaptation culturelle utilisée",
"variant_short":"Version ultra-courte du message (SMS/WhatsApp rapide)"}
Adapte : respect des aînés, relation avant transaction, confiance, Mobile Money..."""
    else:
        return JSONResponse(status_code=400, content={"error": f"Mode inconnu: {req.mode}"})

    messages = [{"role": "system", "content": system}, {"role": "user", "content": req.prompt}]
    try:
        data = await groq_json_with_fallback(messages, max_tokens=3000)
        return JSONResponse(content=data)
    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/projects")
async def projects_ai(req: ProjectRequest):
    from datetime import date
    today_str = date.today().strftime("%d/%m/%Y")

    if req.mode == "planning":
        system = (
            f"Tu es un chef de projet expert PME africaine. Aujourd'hui: {today_str}.\n"
            "Décompose le projet en phases, tâches et délais réalistes, avec contraintes africaines.\n"
            'Retourne UNIQUEMENT un JSON: {"project_name":"...","description":"...","total_duration":"...",'
            '"start_date":"' + today_str + '",'
            '"phases":[{"id":1,"name":"...","duration":"...","color":"#10b981",'
            '"tasks":[{"name":"...","duration":"...","responsible":"...","priority":"haute/moyenne/faible"}],'
            '"milestone":"...","deliverable":"..."}],'
            '"team_needed":["Rôle 1"],"constraints":["Contrainte locale 1"],'
            '"budget_range":"Fourchette FCFA","success_criteria":["Critère 1"]}'
            "\nCouleurs phases: #10b981, #6366f1, #f97316, #3b82f6, #8b5cf6. "
            "Prends en compte: coupures réseau, délais fournisseurs, jours fériés camerounais."
        )

    elif req.mode == "risks":
        system = (
            "Tu es expert gestion des risques PME africaine.\n"
            "Identifie risques et plans de mitigation adaptés au contexte camerounais.\n"
            'Retourne UNIQUEMENT un JSON: {"project_name":"...","overall_risk":"Modéré",'
            '"overall_score":6,'
            '"risks":[{"id":1,"category":"Financier","category_emoji":"💰","risk":"...",'
            '"probability":"Modérée","prob_score":2,"impact":"Élevé","impact_score":3,'
            '"risk_score":6,"mitigation":"Plan proactif","contingency":"Plan B","owner":"Rôle"}],'
            '"recommendations":["Recommandation 1"],"quick_wins":["Action rapide 1"]}'
            "\nEmojis: Financier💰, Technique⚙️, Humain👥, Réglementaire📋, Marché📊, Environnemental🌍"
        )
    else:
        return JSONResponse(status_code=400, content={"error": f"Mode inconnu: {req.mode}"})

    messages = [{"role": "system", "content": system}, {"role": "user", "content": req.prompt}]
    try:
        data = await groq_json_with_fallback(messages, max_tokens=4000)
        return JSONResponse(content=data)
    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/chat")
async def chat(req: ChatRequest):
    result = build_messages_with_rag(req.messages)
    enriched_messages, sources = result if isinstance(result, tuple) else (result, [])

    async def stream():
        try:
            if sources:
                meta = json.dumps({"type": "sources", "data": sources})
                yield f"__META__{meta}__META__"

            async for token in generate(enriched_messages):
                yield token
        except Exception as exc:
            import traceback
            traceback.print_exc()
            print(f"[STREAM ERROR] {type(exc).__name__}: {exc}")
            yield "\n\nUne erreur est survenue. Veuillez réessayer."

    return StreamingResponse(stream(), media_type="text/plain; charset=utf-8")
